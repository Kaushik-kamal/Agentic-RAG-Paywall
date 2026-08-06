"""Document loaders.

Returns *pages* — ``(page_number, text)`` tuples — rather than one blob, so
citations can point a reader at the exact page of the source PDF.
"""

from __future__ import annotations

import csv
import io
import logging
import re
from dataclasses import dataclass
from pathlib import Path

from app.core.errors import ValidationError

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".csv": "text/csv",
    ".json": "application/json",
}


@dataclass(slots=True)
class LoadedPage:
    page: int | None
    text: str


@dataclass(slots=True)
class LoadedDocument:
    title: str
    pages: list[LoadedPage]
    media_type: str

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def char_count(self) -> int:
        return sum(len(p.text) for p in self.pages)

    @property
    def page_count(self) -> int | None:
        numbered = [p.page for p in self.pages if p.page is not None]
        return max(numbered) if numbered else None


_WHITESPACE = re.compile(r"[ \t ]+")
_BLANK_LINES = re.compile(r"\n{3,}")
_HYPHEN_BREAK = re.compile(r"(\w)-\n(\w)")
_CONTROL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def normalise(text: str) -> str:
    """Undo the usual PDF-extraction damage without destroying structure."""
    text = _CONTROL.sub("", text.replace("\r\n", "\n").replace("\r", "\n"))
    text = _HYPHEN_BREAK.sub(r"\1\2", text)  # re-join words split across lines
    text = _WHITESPACE.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return _BLANK_LINES.sub("\n\n", text).strip()


def _load_pdf(path: Path) -> list[LoadedPage]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:
            raise ValidationError(
                "This PDF is password-protected and cannot be indexed."
            ) from exc

    pages: list[LoadedPage] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            raw = page.extract_text() or ""
        except Exception:
            logger.warning("Failed to extract page %d of %s", index, path.name)
            continue
        text = normalise(raw)
        if text:
            pages.append(LoadedPage(page=index, text=text))

    if not pages:
        raise ValidationError(
            "No selectable text found in this PDF. Scanned documents need OCR first."
        )
    return pages


def _load_docx(path: Path) -> list[LoadedPage]:
    import docx  # python-docx

    document = docx.Document(str(path))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "1"
            blocks.append(f"{'#' * min(int(level), 6)} {text}")
        elif style.startswith("list"):
            blocks.append(f"- {text}")
        else:
            blocks.append(text)

    for table in document.tables:
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
        ]
        rows = [r for r in rows if r.replace("|", "").strip()]
        if rows:
            blocks.append("\n".join(rows))

    text = normalise("\n\n".join(blocks))
    if not text:
        raise ValidationError("This DOCX file appears to be empty.")
    return [LoadedPage(page=None, text=text)]


def _load_csv(path: Path) -> list[LoadedPage]:
    raw = path.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(io.StringIO(raw))
    rows = list(reader)
    if not rows:
        raise ValidationError("This CSV file is empty.")

    header, *body = rows
    lines = [" | ".join(header), " | ".join("---" for _ in header)]
    lines.extend(" | ".join(cell.strip() for cell in row) for row in body)
    return [LoadedPage(page=None, text=normalise("\n".join(lines)))]


def _load_text(path: Path) -> list[LoadedPage]:
    text = normalise(path.read_text(encoding="utf-8", errors="replace"))
    if not text:
        raise ValidationError("This file is empty.")
    return [LoadedPage(page=None, text=text)]


def _derive_title(path: Path, pages: list[LoadedPage]) -> str:
    """Prefer a real heading or first meaningful line over the filename."""
    if pages:
        for line in pages[0].text.split("\n")[:6]:
            candidate = line.lstrip("# ").strip()
            if 3 < len(candidate) <= 120 and not candidate.endswith((".", ",", ";")):
                return candidate
    return path.stem.replace("_", " ").replace("-", " ").strip() or path.name


def load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValidationError(
            f"Unsupported file type '{suffix or path.name}'.",
            details={"supported": sorted(SUPPORTED_EXTENSIONS)},
        )

    loader = {
        ".pdf": _load_pdf,
        ".docx": _load_docx,
        ".csv": _load_csv,
    }.get(suffix, _load_text)

    pages = loader(path)
    return LoadedDocument(
        title=_derive_title(path, pages),
        pages=pages,
        media_type=SUPPORTED_EXTENSIONS[suffix],
    )
